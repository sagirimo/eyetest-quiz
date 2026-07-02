#!/usr/bin/env python3
"""Final extractor - handles ALL observed formats in the ophthalmology question bank."""
import docx, pdfplumber, json, re, os

BASE = "/Users/moliex/projects/eyetest"
GROUPS = ["眼科出题第一组", "眼科出题第二组", "眼科出题第三组"]

def extract_docx(fp):
    doc = docx.Document(fp)
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            parts.append(" | ".join(c.text.strip() for c in row.cells))
    return "\n".join(parts)

def extract_pdf(fp):
    texts = []
    with pdfplumber.open(fp) as pdf:
        for page in pdf.pages:
            t = page.extract_text()
            if t: texts.append(t)
    return "\n".join(texts)

def extract_author(fname):
    n = re.sub(r'\.(docx?|pdf)$', '', fname)
    n = re.sub(r'[-_]?\d{10}', '', n)
    n = re.sub(r'[-_]?\d{8,}', '', n)
    n = re.sub(r'[-_()（）\[\]]+', '', n)
    for pat in [r'眼科学?期末.*', r'眼科学?出题.*', r'眼科学?考题.*',
                r'眼科学?试题.*', r'眼科学?作业.*', r'眼科题目.*',
                r'眼科命题.*', r'眼科出题.*', r'眼科考题.*', r'眼科出題.*',
                r'眼科题.*', r'眼科$', r'眼科见习复习题']:
        n = re.sub(pat, '', n)
    return n.strip() if n.strip() and len(n.strip()) >= 2 else fname[:20]

def parse_all(text):
    """Universal parser supporting all observed formats."""
    questions = []
    lines = [l.strip() for l in text.split('\n') if l.strip()]

    # ---- Phase 0: Find answer key block (韦鹏昊-style) ----
    answer_key = {}
    expl_map = {}

    # Look for "答案速查" section
    ak_match = re.search(r'答案速查\s*\n(.+?)(?:\n\s*\n|\n(?:二|三|四)[、\.])', text, re.DOTALL)
    if ak_match:
        for m in re.finditer(r'(\d+)\s*[\.\、]\s*([A-Ea-e])', ak_match.group(1)):
            answer_key[int(m.group(1))] = m.group(2).upper()

    # Look for detailed answer blocks: "N. 【答案】X\n【解析】..."
    for m in re.finditer(r'(\d+)\.\s*【答案】\s*([A-Ea-e])\s*\n【解析】\s*(.+?)(?=\n\d+\.\s*【答案】|\n(?:三|四|五)[、\.]|\Z)', text, re.DOTALL):
        qnum = int(m.group(1))
        ans = m.group(2).upper()
        expl = m.group(3).strip()
        expl_map[qnum] = (ans, expl)

    # ---- Phase 1: Parse questions ----
    i = 0
    current_q = None
    current_opts = {}

    while i < len(lines):
        line = lines[i]

        # Skip separators, headers
        if re.match(r'^-{3,}$', line) or re.match(r'^(?:一|二|三|四|五)[、\.]', line):
            i += 1; continue

        # Pattern 1: 【N】question_text (周智睿, 战鹏宇)
        m1 = re.match(r'^【(\d+)】\s*(.+)', line)
        # Pattern 2: N. question_text or N、question_text
        m2 = re.match(r'^(\d+)\s*[\.\、．]\s*(.+)', line)

        q_match = m1 or m2

        if q_match:
            q_num = int(q_match.group(1))
            q_text = q_match.group(2).strip()

            # Skip if very short or looks like option
            if len(q_text) < 4 or re.match(r'^[A-Ea-e][\.\、\s]', q_text):
                i += 1; continue

            # Save previous question
            if current_q and current_opts and len(current_opts) >= 2:
                current_q['options'] = current_opts
                questions.append(current_q)

            current_q = {'num': q_num, 'question': q_text}
            current_opts = {}
            i += 1; continue

        # Option lines: A. B. C. D. E.  (various separators)
        opt_m = re.match(r'^([A-Ea-e])\s*[\.\、．\s]\s*(.+)', line)
        if opt_m and current_q:
            letter = opt_m.group(1).upper()
            txt = opt_m.group(2).strip()
            if len(txt) >= 1:
                current_opts[letter] = txt
                i += 1; continue

        # Answer lines
        # "答案：X。解析：..." or "答案：X" or "正确答案：X"
        ans_m = re.match(r'^(?:正确)?答案\s*[：:]\s*([A-Ea-e])', line)
        if ans_m and current_q:
            current_q['answer'] = ans_m.group(1).upper()
            # Extract inline explanation
            rest = re.sub(r'^(?:正确)?答案\s*[：:]\s*[A-Ea-e]\s*[。，.]?\s*', '', line)
            rest = re.sub(r'^解析[：:]?\s*', '', rest)
            if rest.strip():
                current_q['explanation'] = rest.strip()
            # Collect continuation explanation lines
            j = i + 1
            extra = []
            while j < len(lines) and j < i + 8:
                nxt = lines[j]
                if re.match(r'^-{3,}$', nxt): j += 1; continue
                if re.match(r'^【(\d+)】', nxt): break
                if re.match(r'^(\d+)\s*[\.\、]', nxt): break
                if re.match(r'^(?:正确)?答案\s*[：:]', nxt): break
                if re.match(r'^(?:一|二|三|四|五)[、\.]', nxt): break
                extra.append(nxt)
                j += 1
            if extra:
                curr = current_q.get('explanation', '')
                current_q['explanation'] = (curr + ' ' + ' '.join(extra)).strip()
            i = j; continue

        # 【答案】X format (邓敏智 style)
        ans_m2 = re.match(r'^【答案】\s*([A-Ea-e])', line)
        if ans_m2 and current_q:
            current_q['answer'] = ans_m2.group(1).upper()
            i += 1; continue

        # Continuation of question text (before options)
        if current_q and not current_opts and len(line) > 5:
            current_q['question'] += ' ' + line
            i += 1; continue

        i += 1

    # Save last question
    if current_q and current_opts and len(current_opts) >= 2:
        current_q['options'] = current_opts
        questions.append(current_q)

    # ---- Phase 2: Apply answers from key/expl maps ----
    for q in questions:
        qnum = q.get('num', 0)
        if q.get('answer'):
            continue
        if qnum in expl_map:
            q['answer'], q['explanation'] = expl_map[qnum]
        elif qnum in answer_key:
            q['answer'] = answer_key[qnum]

    return [q for q in questions if q.get('answer') and len(q.get('options', {})) >= 2]


# ====== MAIN ======
all_qs = []

for group in GROUPS:
    gpath = os.path.join(BASE, group)
    if not os.path.exists(gpath): continue
    for fname in sorted(os.listdir(gpath)):
        if fname.startswith('.'): continue
        fpath = os.path.join(gpath, fname)
        try:
            text = extract_docx(fpath) if fname.endswith('.docx') else extract_pdf(fpath)
        except Exception as e:
            print(f"  SKIP {fname}: {e}")
            continue

        author = extract_author(fname)
        qs = parse_all(text)
        for q in qs:
            q['author'] = author
            q['group'] = group
            q['filename'] = fname

        print(f"{group}/{fname}: {len(qs)}")
        all_qs.extend(qs)

# Dedup
seen = set()
uniq = []
for q in all_qs:
    key = q['question'][:60].strip()
    if key not in seen:
        seen.add(key)
        uniq.append(q)

print(f"\nTotal: {len(all_qs)} raw -> {len(uniq)} unique")

outpath = os.path.join(BASE, 'questions.json')
with open(outpath, 'w', encoding='utf-8') as f:
    json.dump(uniq, f, ensure_ascii=False, indent=2)

authors = {}
for q in uniq:
    a = q['author']
    authors[a] = authors.get(a, 0) + 1
for a, c in sorted(authors.items(), key=lambda x: -x[1]):
    print(f"  {a}: {c}")
print(f"\nSaved {len(uniq)} questions to {outpath}")

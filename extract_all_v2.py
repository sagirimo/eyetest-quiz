#!/usr/bin/env python3
"""Extract all ophthalmology questions - v2 with enhanced parsing."""
import docx, pdfplumber, json, re, os

BASE = "/Users/moliex/projects/eyetest"
GROUPS = ["眼科出题第一组", "眼科出题第二组", "眼科出题第三组"]


def extract_docx(filepath):
    doc = docx.Document(filepath)
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            parts.append(" | ".join(cells))
    return "\n".join(parts)


def extract_pdf(filepath):
    texts = []
    with pdfplumber.open(filepath) as pdf:
        for page in pdf.pages:
            t = page.extract_text()
            if t:
                texts.append(t)
    return "\n".join(texts)


def extract_author(filename, group):
    name = filename
    name = re.sub(r'\.(docx?|pdf)$', '', name)
    name = re.sub(r'[-_]?\d{10}', '', name)
    name = re.sub(r'[-_]?\d{8,}', '', name)
    name = re.sub(r'[-_()（）\[\]]+', '', name)
    name = re.sub(r'眼科学?期末.*', '', name)
    name = re.sub(r'眼科学?出题.*', '', name)
    name = re.sub(r'眼科学?考题.*', '', name)
    name = re.sub(r'眼科学?试题.*', '', name)
    name = re.sub(r'眼科学?作业.*', '', name)
    name = re.sub(r'眼科题目.*', '', name)
    name = re.sub(r'眼科命题.*', '', name)
    name = re.sub(r'眼科出题.*', '', name)
    name = re.sub(r'眼科考题.*', '', name)
    name = re.sub(r'眼科出題.*', '', name)
    name = re.sub(r'眼科题.*', '', name)
    name = re.sub(r'眼科$', '', name)
    name = re.sub(r'眼科见习复习题', '', name)
    if not name or len(name) < 2:
        return filename[:20]
    return name.strip()


def parse_answer_key_block(text_block):
    """Parse answer key like '1.B　　2.C　　3.B' into dict."""
    answers = {}
    # Pattern: number.letter with various spacing
    for m in re.finditer(r'(\d+)\s*[\.\、]\s*([A-Ea-e])', text_block):
        answers[int(m.group(1))] = m.group(2).upper()
    return answers


def parse_questions_v2(text, author, group, filename):
    """Enhanced question parser that handles multiple formats."""
    questions = []
    lines = text.split('\n')
    lines = [l.strip() for l in lines]

    # STEP 1: Find answer key block (if exists)
    answer_key = {}
    ans_key_match = re.search(
        r'(?:答案速查|参考答案|选择题答案).*?\n(.*?)(?:\n\n|\n(?:二|三|四|五)[、\.])',
        text, re.DOTALL)
    if ans_key_match:
        answer_key = parse_answer_key_block(ans_key_match.group(1))
        # Also try line-by-line
        for line in lines:
            if re.match(r'^\d+[\.\、]\s*[A-Ea-e](?:\s|$)', line):
                m = re.match(r'(\d+)[\.\、]\s*([A-Ea-e])', line)
                answer_key[int(m.group(1))] = m.group(2).upper()

    # STEP 2: Scan for detailed answer blocks (like 【答案】X or 答案：X)
    detailed_answers = {}  # q_num -> (answer, explanation)
    i = 0
    while i < len(lines):
        # Pattern: 【数字】【答案】X or 数字.【答案】X
        m = re.match(r'(?:【)?(\d+)(?:】)?[\.\、\s]*【?答案】?\s*[：:]?\s*([A-Ea-e])', lines[i])
        if not m:
            m = re.match(r'^答案\s*[：:]\s*([A-Ea-e])', lines[i])

        if m:
            try:
                q_num = int(m.group(1))
                ans = m.group(2).upper()
            except (ValueError, IndexError):
                i += 1
                continue
            # Collect explanation from following lines
            expl_lines = []
            j = i + 1
            while j < len(lines) and j < i + 10:
                nxt = lines[j]
                if re.match(r'(?:【)?(\d+)(?:】)?[\.\、\s]*【?答案】?', nxt):
                    break
                if re.match(r'^(?:【)?(\d+)(?:】)?[\.\、]', nxt) and '答案' not in nxt:
                    break  # Next question
                expl_text = nxt
                expl_text = re.sub(r'^【?解析】?\s*[：:]?\s*', '', expl_text)
                if expl_text:
                    expl_lines.append(expl_text)
                j += 1
            detailed_answers[q_num] = (ans, ' '.join(expl_lines))
            i = j
            continue
        i += 1

    # STEP 3: Extract questions
    i = 0
    current_q = None
    current_opts = {}
    in_opts = False
    last_q_num = 0

    while i < len(lines):
        line = lines[i]

        # Detect question start: 【N】or N. or N、
        q_match = re.match(r'^【?(\d+)】?\s*[\.\、．]\s*(.+)', line)
        if not q_match:
            q_match = re.match(r'^(\d+)\s*[\.\、．]\s*(.+)', line)

        if q_match:
            q_num = int(q_match.group(1))
            q_text = q_match.group(2).strip()

            # Filter false positives: if "q_text" is actually an option
            if re.match(r'^[A-Ea-e][\.\、\s]', q_text):
                i += 1
                continue

            # Save previous
            if current_q and current_opts and len(current_opts) >= 2:
                current_q['options'] = current_opts
                questions.append(current_q)

            current_q = {
                'num': q_num,
                'question': q_text,
                'author': author,
                'group': group,
                'filename': filename,
                'topic': '',
            }
            current_opts = {}
            in_opts = True
            last_q_num = q_num
            i += 1
            continue

        # Detect option lines: A. B. C. D. E.
        opt_match = re.match(r'^([A-Ea-e])\s*[\.\、．]\s*(.+)', line)
        if opt_match and in_opts and current_q:
            opt_letter = opt_match.group(1).upper()
            opt_text = opt_match.group(2).strip()
            if len(opt_text) >= 2:
                current_opts[opt_letter] = opt_text
                i += 1
                continue

        # Detect answer marker for current question
        ans_marker = re.match(
            r'^(?:【)?答案(?:】)?\s*[：:]\s*([A-Ea-e])\s*(?:[。，]?\s*解析[：:]?(.*))?', line)
        if ans_marker and current_q:
            current_q['answer'] = ans_marker.group(1).upper()
            expl = ans_marker.group(2) if ans_marker.lastindex >= 2 else ''
            if expl:
                current_q['explanation'] = expl.strip()
            # Check for more explanation on next lines
            j = i + 1
            extra = []
            while j < len(lines) and j < i + 5:
                nxt = lines[j]
                if re.match(r'^(?:【)?(\d+)(?:】)?[\.\、]', nxt):
                    break
                if re.match(r'^(?:【)?答案', nxt):
                    break
                extra.append(nxt)
                j += 1
            if extra:
                current_q['explanation'] = (current_q.get('explanation', '') + ' ' + ' '.join(extra)).strip()
            i = j
            continue

        # Detect 正确答案 marker
        ans_marker2 = re.match(
            r'^正确答案\s*[：:]\s*([A-Ea-e])\s*(?:[。，]?\s*解析[：:]?(.*))?', line)
        if ans_marker2 and current_q:
            current_q['answer'] = ans_marker2.group(1).upper()
            expl = ans_marker2.group(2) if ans_marker2.lastindex >= 2 else ''
            if expl:
                current_q['explanation'] = expl.strip()
            i += 1
            continue

        # If collecting options and line doesn't match anything
        if in_opts and current_q and not current_opts:
            # Might be continuation of question
            if len(line) > 5 and not line.startswith('【'):
                current_q['question'] += ' ' + line
            i += 1
            continue

        i += 1

    # Save last question
    if current_q and current_opts and len(current_opts) >= 2:
        current_q['options'] = current_opts
        questions.append(current_q)

    # STEP 4: Apply answers from answer_key or detailed_answers
    for q in questions:
        qnum = q.get('num', 0)

        # Prioritize inline answer (already set during parsing)
        if q.get('answer'):
            continue

        # Check detailed answers
        if qnum in detailed_answers:
            ans, expl = detailed_answers[qnum]
            q['answer'] = ans
            if expl:
                q['explanation'] = expl
            continue

        # Check answer key
        if qnum in answer_key:
            q['answer'] = answer_key[qnum]
            continue

    # Filter valid questions
    valid = [q for q in questions if q.get('answer') and len(q.get('options', {})) >= 2]
    return valid


def parse_pdf_questions(text, author, group, filename):
    """Special parser for PDF format (孙晟佳 style - answers inline in 【】)."""
    questions = []
    lines = text.split('\n')
    lines = [l.strip() for l in lines if l.strip()]

    i = 0
    current_q = None
    current_opts = {}
    shared_options = {}  # For "共用选项" type questions

    while i < len(lines):
        line = lines[i]

        # Check for shared options header
        shared_match = re.match(r'^(\d+)～(\d+)题共用选项', line)
        if shared_match:
            # Read shared options from next lines
            j = i + 1
            shared_options = {}
            while j < len(lines) and j < i + 10:
                opt_line = lines[j]
                opt_m = re.match(r'^([A-E])\s+(.+)', opt_line)
                if opt_m:
                    shared_options[opt_m.group(1)] = opt_m.group(2).strip()
                    j += 1
                    continue
                # If we hit a question number, stop
                if re.match(r'^\d+\s', opt_line):
                    break
                j += 1
            i = j
            continue

        # Question start: "N question_text"
        q_match = re.match(r'^(\d+)\s+(.+)', line)
        if q_match and not re.match(r'^\d+[\.\、]', line):
            q_num = int(q_match.group(1))
            q_text = q_match.group(2).strip()

            # Save previous
            if current_q and current_opts and len(current_opts) >= 2:
                current_q['options'] = current_opts
                questions.append(current_q)

            current_q = {
                'num': q_num,
                'question': q_text,
                'author': author,
                'group': group,
                'filename': filename,
                'topic': '',
            }
            # Use shared options if available
            current_opts = dict(shared_options) if shared_options else {}
            i += 1
            continue

        # Option lines
        opt_match = re.match(r'^([A-E])\s+(.+)', line)
        if opt_match and current_q and not shared_options:
            opt_letter = opt_match.group(1)
            opt_text = opt_match.group(2).strip()
            if len(opt_text) >= 2:
                current_opts[opt_letter] = opt_text
                i += 1
                continue

        # Inline answer: 【X】
        ans_match = re.search(r'【([A-Ea-e])】', line)
        if ans_match and current_q:
            current_q['answer'] = ans_match.group(1).upper()
            # The question text might include the answer marker - clean it
            current_q['question'] = re.sub(r'【[A-Ea-e]】', '', current_q.get('question', '')).strip()
            i += 1
            continue

        i += 1

    # Save last
    if current_q and current_opts and len(current_opts) >= 2:
        current_q['options'] = current_opts
        questions.append(current_q)

    # Filter
    valid = [q for q in questions if q.get('answer') and len(q.get('options', {})) >= 2]
    return valid


def main():
    all_questions = []

    for group in GROUPS:
        group_path = os.path.join(BASE, group)
        if not os.path.exists(group_path):
            continue

        for filename in sorted(os.listdir(group_path)):
            if filename.startswith('.'):
                continue

            filepath = os.path.join(group_path, filename)
            print(f"Processing: {group}/{filename}")

            try:
                if filename.endswith('.docx'):
                    text = extract_docx(filepath)
                elif filename.endswith('.pdf'):
                    text = extract_pdf(filepath)
                else:
                    continue
            except Exception as e:
                print(f"  ERROR extracting: {e}")
                continue

            author = extract_author(filename, group)

            # Try v2 parser
            mcqs = parse_questions_v2(text, author, group, filename)

            # If PDF and few results, try PDF-specific parser
            if filename.endswith('.pdf') and len(mcqs) < 10:
                pdf_mcqs = parse_pdf_questions(text, author, group, filename)
                if len(pdf_mcqs) > len(mcqs):
                    mcqs = pdf_mcqs

            print(f"  Found {len(mcqs)} MCQs")
            all_questions.extend(mcqs)

    # Deduplicate
    print(f"\nTotal before dedup: {len(all_questions)}")
    seen = set()
    deduped = []
    for q in all_questions:
        key = q['question'][:60].strip()
        if key not in seen:
            seen.add(key)
            deduped.append(q)

    print(f"Total after dedup: {len(deduped)}")

    # Save
    outpath = os.path.join(BASE, 'questions.json')
    with open(outpath, 'w', encoding='utf-8') as f:
        json.dump(deduped, f, ensure_ascii=False, indent=2)
    print(f"Saved to {outpath}")

    # Stats
    authors = {}
    for q in deduped:
        a = q['author']
        authors[a] = authors.get(a, 0) + 1
    print("\nQuestions by author:")
    for a, c in sorted(authors.items(), key=lambda x: -x[1]):
        print(f"  {a}: {c}")


if __name__ == '__main__':
    main()

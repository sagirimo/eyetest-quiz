#!/usr/bin/env python3
"""V5 — definitive extractor. Every format, every file."""
import docx, pdfplumber, json, re, os

BASE = "/Users/moliex/projects/eyetest"
GROUPS = ["眼科出题第一组", "眼科出题第二组", "眼科出题第三组"]

def extract_docx(fp):
    doc = docx.Document(fp)
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            parts.append(" | ".join(c.text.strip() for c in row.cells))
    return "\n".join(parts)

def extract_pdf(fp):
    return "\n".join([p.extract_text() or '' for p in pdfplumber.open(fp).pages])

def clean_author(fname):
    n = re.sub(r'\.(docx?|pdf)$', '', fname)
    n = re.sub(r'[-_]?\d{10}', '', n)
    n = re.sub(r'[-_]?\d{8,}', '', n)
    n = re.sub(r'[-_()（）\[\]]+', '', n)
    for pat in [r'眼科学?期末.*', r'眼科学?出题.*', r'眼科学?考题.*',
                r'眼科学?试题.*', r'眼科学?作业.*', r'眼科题目.*',
                r'眼科命题.*', r'眼科出题.*', r'眼科考题.*', r'眼科出題.*',
                r'眼科题.*', r'眼科$', r'眼科见习复习题']:
        n = re.sub(pat, '', n)
    return n.strip()[:30] if n.strip() else fname[:20]

AUTHOR_FIXES = {
    '谭饮骏-2110122901':'谭饮骏','韦鹏昊眼科学综合试题':'韦鹏昊','刘霁影45题1':'刘霁影',
    '眼科出题-王雨涵.docx':'王雨涵','眼科考题刘子越.docx':'刘子越',
    '眼科期末出题-陈栐洁2210902818':'陈栐洁',
    '眼科作业-2210122920-梁颖欣.':'梁颖欣',
    '眼科作业 梁瑞博(5).docx':'梁瑞博',
    '眼科出題_陳揆文.docx':'陳揆文',
}

def parse(text, fname):
    """Parse with format auto-detection."""
    lines = [l.strip() for l in text.split('\n') if l.strip()]

    # ═══ Determine primary format ═══
    has_inline_bracket = any(re.search(r'[\[【]([A-Ea-e])[\]】]', l) for l in lines)
    has_answer_per_line = any(re.match(r'^答案[：:]', l) for l in lines)
    has_daan_marker = any(re.match(r'^(?:【?正确】?)?(?:正确)?答案', l) for l in lines)
    has_trailing_ans = any(re.search(r'[？?）\)]\s*[A-Ea-e]\s*$', l) for l in lines)
    has_trailing_ans2 = any(re.search(r'\s{2,}[A-Ea-e]\s*$', l) for l in lines)
    has_diNti = any(re.match(r'^第\s*\d+\s*题', l) for l in lines)
    has_seq_unanchored = any(re.match(r'^(\d+)\s*[\.\、]', l) for l in lines)
    has_ordered_ref = '参考答案' in text

    questions = []

    # ═══ PHASE 0: Answer sources ═══
    # Source A: Answer key block (韦鹏昊)
    key_from_block = {}
    ak = re.search(r'答案速查\s*\n(.+?)(?:\n\s*\n|\n(?:二|三|四)[、\.])', text, re.DOTALL)
    if ak:
        for m in re.finditer(r'(\d+)\s*[\.\、]\s*([A-Ea-e])', ak.group(1)):
            key_from_block[int(m.group(1))] = m.group(2).upper()

    # Source B: Detailed answer blocks
    expl_map = {}
    for m in re.finditer(
        r'(\d+)\.\s*【答案】\s*([A-Ea-e])\s*\n【解析】\s*(.+?)(?=\n\d+\.\s*【答案】|\n(?:三|四|五)[、\.]|\Z)',
        text, re.DOTALL):
        expl_map[int(m.group(1))] = (m.group(2).upper(), m.group(3).strip())

    # Source C: 陈栐洁 ordered answers (letters only, no numbers)
    ordered_answers = []
    if has_ordered_ref:
        ref = re.search(r'参考答案\s*\n一、单选题[^\n]*\n(.+?)(?:\n\s*\n|\n(?:二|三)[、\.\n])', text, re.DOTALL)
        if ref:
            for line in ref.group(1).split('\n'):
                m = re.match(r'^\s*([A-Ea-e])\s*[（(]', line) or re.match(r'^\s*([A-Ea-e])\s*$', line)
                if m: ordered_answers.append(m.group(1).upper())

    # ═══ PHASE 1: Parse questions ═══
    i = 0
    current_q = None
    current_opts = {}
    seq_num = 0
    Q = []  # local accumulator for questions

    def save_q():
        nonlocal current_q, current_opts
        if current_q and current_opts and len(current_opts) >= 2:
            current_q['options'] = dict(current_opts)
            Q.append(current_q)
        current_q = None
        current_opts = {}

    while i < len(lines):
        line = lines[i]

        # Skip noise
        if (re.match(r'^-{3,}$', line) or
            re.match(r'^(?:一|二|三|四|五|六|七|八|九|十)[、\.\s]', line) or
            re.match(r'^【第[一二三四五六七八九十]', line) or
            re.match(r'^【[^】\d]+】$', line) or  # 【绪论与解剖生理】
            line in ('（单选）', '(单选)', '选择题', '单选题')):
            i += 1; continue

        # ── Detect question start ──
        q_num = None; q_text = None; skip = 0

        # A: 【N】text
        m = re.match(r'^【(\d+)】\s*(.+)', line)
        if m and len(m.group(2)) >= 2:
            q_num, q_text = int(m.group(1)), m.group(2)

        # B: N. text / N、text
        if not q_num:
            m = re.match(r'^(\d+)\s*[\.\、]\s*(.+)', line)
            if m and len(m.group(2)) >= 2:
                q_num = int(m.group(1))
                qt = m.group(2)
                # Remove trailing [X] or 【X】 answer markers
                qt = re.sub(r'\s*[\[【]([A-Ea-e])[\]】]\s*$', '', qt)
                # Remove trailing single-space letter answer (张智翔, 宋娟范, 罗智谦, 曾宇桐)
                # But only if it looks like a real answer (preceded by ? or ）etc)
                am = re.search(r'[？?）\)]\s*([A-Ea-e])\s*$', qt)
                if am:
                    current_trailing = am.group(1).upper()
                    qt = re.sub(r'[？?）\)]\s*[A-Ea-e]\s*$', '？', qt)
                else:
                    am = re.search(r'\s{2,}([A-Ea-e])\s*$', qt)
                    if am:
                        current_trailing = am.group(1).upper()
                        qt = re.sub(r'\s{2,}[A-Ea-e]\s*$', '', qt)
                    else:
                        current_trailing = None
                q_text = qt

        # C: 第 N 题 (text on next line)
        if not q_num:
            m = re.match(r'^第\s*(\d+)\s*题\s*$', line)
            if m and i+1 < len(lines):
                q_num, q_text = int(m.group(1)), lines[i+1]
                skip = 1

        # D: 陈宣仁 style — inline options on same line, answer on next
        # "问题？ A. xxx B. xxx C. xxx D. xxx"
        if not q_num and has_answer_per_line:
            # Detect question+options on one line
            m = re.match(r'^(.+[？?])\s+(A\..+B\..+C\..+D\..+)', line)
            if m:
                seq_num += 1
                q_num = seq_num
                q_text = m.group(1)
                # Parse inline options
                opts_str = m.group(2)
                for om in re.finditer(r'([A-Ea-e])\.\s*(.+?)(?=\s+[A-Ea-e]\.|$)', opts_str):
                    current_opts[om.group(1).upper()] = om.group(2).strip()
                # Answer on next line: "答案：X"
                if i+1 < len(lines):
                    am = re.match(r'^答案[：:]\s*([A-Ea-e])', lines[i+1])
                    if am:
                        save_q()
                        Q[-1]['answer'] = am.group(1).upper()
                        Q[-1]['num'] = q_num
                        i += 2; continue
                i += 1; continue

        # E: Unnumbered question (邓敏智, 陈栐洁, 梁瑞博, 宋娟范 first few)
        if not q_num and not current_q:
            is_opt = bool(re.match(r'^[A-Ea-e]\s*[\.\、\s]', line))
            is_ans = bool(re.match(r'^(?:【?正确】?)?(?:正确)?答案', line))
            is_chapter = bool(re.match(r'^【', line))
            if not is_opt and not is_ans and not is_chapter and len(line) > 8:
                seq_num += 1
                q_num = seq_num
                q_text = line
                # Detect trailing answer
                am = re.search(r'[？?）\)]([A-Ea-e])\s*$', q_text)
                if am:
                    current_trailing = am.group(1).upper()
                    q_text = re.sub(r'[？?）\)][A-Ea-e]\s*$', '？', q_text)
                else:
                    am = re.search(r'\s{2,}([A-Ea-e])\s*$', q_text)
                    if am:
                        current_trailing = am.group(1).upper()
                        q_text = re.sub(r'\s{2,}[A-Ea-e]\s*$', '', q_text)
                    else:
                        current_trailing = None

        if q_num and q_text:
            save_q()
            current_q = {'num': q_num, 'question': q_text.strip()}
            current_opts = {}
            if 'current_trailing' in dir() and current_trailing:
                current_q['_ia'] = current_trailing
                current_trailing = None
            i += 1 + skip; continue

        # ── Option lines ──
        opt_m = re.match(r'^([A-Ea-e])\s*[\.\、．\s]\s*(.+)', line)
        if opt_m and current_q:
            ltr, txt = opt_m.group(1).upper(), opt_m.group(2).strip()
            if len(txt) >= 1:
                current_opts[ltr] = txt
                i += 1; continue

        # ── Answer markers ──
        ans_m = (re.match(r'^(?:【?正确】?)?(?:正确)?答案[：:】]\s*([A-Ea-e])', line) or
                 re.match(r'^【答案】\s*([A-Ea-e])', line))
        if ans_m and current_q:
            current_q['answer'] = ans_m.group(1).upper()
            rest = re.sub(r'^(?:【?正确】?)?(?:正确)?答案[：:】]\s*[A-Ea-e]\s*[。，.]?\s*', '', line)
            rest = re.sub(r'^解析[：:]?\s*', '', rest)
            if rest.strip(): current_q['explanation'] = rest.strip()
            j = i + 1; extra = []
            while j < len(lines) and j < i + 10:
                nxt = lines[j]
                if re.match(r'^-{3,}$', nxt): j += 1; continue
                if any(re.match(p, nxt) for p in [
                    r'^【(\d+)】', r'^(\d+)\s*[\.\、]', r'^(?:【?正确】?)?(?:正确)?答案',
                    r'^(?:一|二|三|四|五)[、\.]', r'^第\s*\d+\s*题']):
                    break
                expl_line = re.sub(r'^【解析】\s*', '', nxt)
                if expl_line: extra.append(expl_line)
                j += 1
            if extra:
                curr = current_q.get('explanation', '')
                current_q['explanation'] = (curr + ' ' + ' '.join(extra)).strip()
            i = j; continue

        # ── 陈宣仁: "答案：X" on its own line ──
        if has_answer_per_line:
            am = re.match(r'^答案[：:]\s*([A-Ea-e])', line)
            if am and current_q:
                current_q['answer'] = am.group(1).upper()
                i += 1; continue

        # ── Question text continuation ──
        if current_q and not current_opts and len(line) > 3:
            current_q['question'] += ' ' + line
            i += 1; continue

        i += 1

    save_q()

    # ═══ PHASE 2: Apply answers ═══
    # For 陈栐洁: ordered answers assigned sequentially
    if ordered_answers:
        for idx, q in enumerate(Q):
            if not q.get('answer') and idx < len(ordered_answers):
                q['answer'] = ordered_answers[idx]

    for q in Q:
        qnum = q.get('num', 0)
        if not q.get('answer') and q.get('_ia'):
            q['answer'] = q.pop('_ia')
        if not q.get('answer') and qnum in key_from_block:
            q['answer'] = key_from_block[qnum]
        if not q.get('answer') and qnum in expl_map:
            q['answer'], q['explanation'] = expl_map[qnum]
        q.pop('_ia', None)

    return [q for q in Q if q.get('answer') and len(q.get('options', {})) >= 2]


# ═══ MAIN ═══
all_qs = []
for group in GROUPS:
    gpath = os.path.join(BASE, group)
    if not os.path.exists(gpath): continue
    for fname in sorted(os.listdir(gpath)):
        if fname.startswith('.'): continue
        fpath = os.path.join(gpath, fname)
        try:
            text = extract_docx(fpath) if fname.endswith('.docx') else extract_pdf(fpath)
        except Exception as e:
            print(f"  SKIP {fname}: {e}"); continue

        author = AUTHOR_FIXES.get(clean_author(fname), clean_author(fname))
        qs = parse(text, fname)
        for q in qs:
            q['author'] = author; q['group'] = group; q['filename'] = fname
        print(f"{fname:<45s} -> {len(qs):3d}  [{author}]")
        all_qs.extend(qs)

# Dedup
seen = set(); uniq = []
for q in all_qs:
    key = q['question'][:80].strip()
    if key not in seen: seen.add(key); uniq.append(q)

print(f"\n{'='*60}")
print(f"Total: {len(all_qs)} raw -> {len(uniq)} unique")

outpath = os.path.join(BASE, 'questions.json')
with open(outpath, 'w', encoding='utf-8') as f:
    json.dump(uniq, f, ensure_ascii=False, indent=2)

from collections import Counter
for a, c in Counter(q['author'] for q in uniq).most_common():
    print(f"  {a}: {c}")
print(f"Saved to {outpath}")

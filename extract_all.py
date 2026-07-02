#!/usr/bin/env python3
"""Extract all ophthalmology questions into structured JSON."""
import docx
import pdfplumber
import json
import re
import os
import sys

BASE = "/Users/moliex/projects/eyetest"
GROUPS = ["眼科出题第一组", "眼科出题第二组", "眼科出题第三组"]


def extract_docx(filepath):
    """Extract text from docx, including paragraphs and tables."""
    doc = docx.Document(filepath)
    parts = []
    for p in doc.paragraphs:
        parts.append(p.text)
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            parts.append(" | ".join(cells))
    return "\n".join(parts)


def extract_pdf(filepath):
    """Extract text from PDF."""
    texts = []
    with pdfplumber.open(filepath) as pdf:
        for page in pdf.pages:
            t = page.extract_text()
            if t:
                texts.append(t)
    return "\n".join(texts)


def extract_author(filename, group):
    """Extract author name from filename."""
    # Patterns like "战鹏宇眼科学..." or "邬清杨-2210122916-眼科..."
    name = filename
    # Remove extension
    name = re.sub(r'\.(docx?|pdf)$', '', name)
    # Remove student IDs
    name = re.sub(r'[-_]?\d{10}', '', name)
    name = re.sub(r'[-_]?\d{8,}', '', name)
    # Remove common suffixes
    name = re.sub(r'[-_()（）]', '', name)
    name = re.sub(r'眼科学?期末.*', '', name)
    name = re.sub(r'眼科学?出题.*', '', name)
    name = re.sub(r'眼科学?考题.*', '', name)
    name = re.sub(r'眼科学?试题.*', '', name)
    name = re.sub(r'眼科学?作业.*', '', name)
    name = re.sub(r'眼科题目.*', '', name)
    name = re.sub(r'眼科命题.*', '', name)
    name = re.sub(r'眼科出题.*', '', name)
    name = re.sub(r'眼科考题.*', '', name)
    name = re.sub(r'眼科出題.*', '', name)
    name = re.sub(r'眼科$', '', name)
    # Also handle "眼科见习复习题" pattern
    name = re.sub(r'眼科见习复习题', '', name)
    if not name or len(name) < 2:
        return filename[:20]
    return name.strip()


def parse_questions_from_text(text, author, group, filename):
    """Parse MCQs from extracted text.

    Returns list of {question, options: {A,B,C,D,E}, answer, explanation, author, group, topic}
    """
    questions = []
    lines = text.split('\n')

    # Clean lines - remove excessive whitespace but keep structure
    lines = [l.strip() for l in lines if l.strip()]

    # Strategy 1: Find question blocks with explicit "答案" markers
    # Pattern: question number -> question text -> options -> 答案：X -> optional 解析

    # First pass: find all question start patterns
    q_patterns = [
        r'^【(\d+)】(.+)',  # 【1】question
        r'^(\d+)[\.\、．]\s*(.+)',  # 1. question or 1、question
        r'^第\s*(\d+)\s*题[：:]\s*(.+)',  # 第1题：question
    ]

    # Find answer patterns
    ans_patterns = [
        r'^答案[：:]\s*([A-Ea-e])',
        r'^正确答案[：:]\s*([A-Ea-e])',
        r'^【答案】[：:]\s*([A-Ea-e])',
        r'^答案[：:]\s*([A-Ea-e])[。，]?\s*解析[：:]?(.*)',
        r'^正确答案[：:]\s*([A-Ea-e])[。，]?\s*解析[：:]?(.*)',
        r'【答案】\s*([A-Ea-e])',
    ]

    # Build a combined scan: walk through lines and build question blocks
    i = 0
    current_q = None
    current_options = {}
    in_options = False

    while i < len(lines):
        line = lines[i]

        # Check if this line starts a numbered question
        q_match = None
        q_num = None

        for pat in q_patterns:
            m = re.match(pat, line)
            if m:
                q_match = m
                q_num = int(m.group(1))
                break

        if q_match and q_num:
            # Save previous question if exists
            if current_q and current_q.get('question') and current_options:
                current_q['options'] = current_options
                questions.append(current_q)

            # Start new question - the question text may continue on next lines
            q_text = q_match.group(2).strip()
            current_q = {
                'question': q_text,
                'author': author,
                'group': group,
                'filename': filename,
                'topic': '',
            }
            current_options = {}
            in_options = True
            i += 1
            continue

        # Check for answer line
        ans_match = None
        for pat in ans_patterns:
            m = re.match(pat, line)
            if m:
                ans_match = m
                break

        if ans_match and current_q:
            current_q['answer'] = ans_match.group(1).upper()
            explanation = ''
            if len(ans_match.groups()) > 1 and ans_match.group(2):
                explanation = ans_match.group(2).strip()
            # Check next lines for more explanation
            j = i + 1
            while j < len(lines) and j < i + 5:
                next_line = lines[j]
                # Stop if next line looks like a new question
                is_new_q = False
                for pat in q_patterns:
                    if re.match(pat, next_line):
                        is_new_q = True
                        break
                if is_new_q or re.match(r'^答案[：:]', next_line) or re.match(r'^正确答案[：:]', next_line):
                    break
                if next_line.startswith('解析') or next_line.startswith('【解析】'):
                    explanation += ' ' + re.sub(r'^【?解析[：:]?】?', '', next_line).strip()
                elif explanation:
                    explanation += ' ' + next_line
                j += 1
            current_q['explanation'] = explanation.strip()
            i += 1
            continue

        # Check for option lines (A. B. C. D. E.)
        option_match = re.match(r'^([A-Ea-e])[\.\、．\s]+(.+)', line)
        if option_match and in_options and current_q:
            opt_letter = option_match.group(1).upper()
            opt_text = option_match.group(2).strip()
            # Don't capture if it looks like a regular sentence starting with A.
            if len(opt_text) > 2:  # Real option text
                current_options[opt_letter] = opt_text
                i += 1
                continue

        # If we're collecting options and see continuation text
        if in_options and current_q and not current_options:
            # Might be continuation of question text
            current_q['question'] += ' ' + line
            i += 1
            continue

        i += 1

    # Save last question
    if current_q and current_q.get('question') and current_options:
        current_q['options'] = current_options
        questions.append(current_q)

    # Post-process: filter out questions without answers or with too few options
    valid = []
    for q in questions:
        if q.get('answer') and len(q.get('options', {})) >= 2:
            valid.append(q)

    return valid


def parse_long_form_questions(text, author, group, filename):
    """Parse long-form questions (简答题, 病例分析题, etc.)"""
    long_questions = []
    lines = text.split('\n')
    lines = [l.strip() for l in lines if l.strip()]

    # Patterns for long-form question starts
    lq_starts = [
        r'^(\d+)[\.\、．]\s*简答题[：:]?\s*(.+)',
        r'^(\d+)[\.\、．]\s*病例分析题?[：:]?\s*(.+)',
        r'^简答题[：:]\s*(.+)',
        r'^病例分析[：:]\s*(.+)',
        r'^(\d+)[\.\、．]\s*(简述.+)',
    ]

    answer_markers = ['答：', '答案：', '参考答案：', '【答案】', '【解析】']

    i = 0
    while i < len(lines):
        line = lines[i]

        is_lq = False
        for pat in lq_starts:
            m = re.match(pat, line)
            if m:
                is_lq = True
                break

        if is_lq:
            q_text = line
            # Collect question text
            j = i + 1
            while j < len(lines) and j < i + 10:
                next_line = lines[j]
                is_ans = any(next_line.startswith(am) for am in answer_markers)
                is_new_q = False
                for pat in lq_starts:
                    if re.match(pat, next_line):
                        is_new_q = True
                        break
                if is_ans or is_new_q:
                    break
                q_text += '\n' + next_line
                j += 1

            # Look for answer
            answer = ''
            k = j
            while k < len(lines) and k < j + 30:
                aline = lines[k]
                is_new_q = False
                for pat in lq_starts:
                    if re.match(pat, aline):
                        is_new_q = True
                        break
                if is_new_q and k > j:
                    break
                for am in answer_markers:
                    if aline.startswith(am):
                        answer += aline[len(am):] + '\n'
                        break
                else:
                    if answer:  # Continue collecting answer
                        answer += aline + '\n'
                k += 1

            long_questions.append({
                'question': q_text,
                'answer': answer.strip(),
                'author': author,
                'group': group,
                'filename': filename,
                'type': 'long',
            })
            i = k
            continue
        i += 1

    return long_questions


def main():
    all_questions = []
    all_long_questions = []

    for group in GROUPS:
        group_path = os.path.join(BASE, group)
        if not os.path.exists(group_path):
            continue

        for filename in sorted(os.listdir(group_path)):
            if filename.startswith('.'):
                continue

            filepath = os.path.join(group_path, filename)
            print(f"Processing: {group}/{filename}")

            try:
                if filename.endswith('.docx'):
                    text = extract_docx(filepath)
                elif filename.endswith('.pdf'):
                    text = extract_pdf(filepath)
                else:
                    continue
            except Exception as e:
                print(f"  ERROR extracting: {e}")
                continue

            author = extract_author(filename, group)

            # Parse MCQs
            mcqs = parse_questions_from_text(text, author, group, filename)
            print(f"  Found {len(mcqs)} MCQs")
            all_questions.extend(mcqs)

            # Parse long-form
            long_qs = parse_long_form_questions(text, author, group, filename)
            print(f"  Found {len(long_qs)} long-form questions")
            all_long_questions.extend(long_qs)

    # Deduplicate questions by similarity
    print(f"\nTotal MCQs before dedup: {len(all_questions)}")

    # Simple dedup: same first 30 chars of question
    seen = set()
    deduped = []
    for q in all_questions:
        key = q['question'][:50].strip()
        if key not in seen:
            seen.add(key)
            deduped.append(q)

    print(f"Total MCQs after dedup: {len(deduped)}")
    print(f"Total long-form: {len(all_long_questions)}")

    # Save
    output = {
        'mcq': deduped,
        'long_form': all_long_questions,
    }

    outpath = os.path.join(BASE, 'questions.json')
    with open(outpath, 'w', encoding='utf-8') as f:
        json.dump(output, f, ensure_ascii=False, indent=2)

    print(f"\nSaved to {outpath}")

    # Print stats by author
    authors = {}
    for q in deduped:
        a = q['author']
        if a not in authors:
            authors[a] = 0
        authors[a] += 1

    print("\nQuestions by author:")
    for a, count in sorted(authors.items(), key=lambda x: -x[1]):
        print(f"  {a}: {count}")


if __name__ == '__main__':
    main()

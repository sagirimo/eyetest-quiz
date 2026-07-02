#!/usr/bin/env python3
"""V3 extractor - handles ALL observed formats across all three groups."""
import docx, pdfplumber, json, re, os

BASE = "/Users/moliex/projects/eyetest"
GROUPS = ["眼科出题第一组", "眼科出题第二组", "眼科出题第三组"]

def extract_docx(fp):
    doc = docx.Document(fp)
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            parts.append(" | ".join(c.text.strip() for c in row.cells))
    return "\n".join(parts)

def extract_pdf(fp):
    texts = []
    with pdfplumber.open(fp) as pdf:
        for page in pdf.pages:
            t = page.extract_text()
            if t: texts.append(t)
    return "\n".join(texts)

def extract_author(fname):
    n = re.sub(r'\.(docx?|pdf)$', '', fname)
    n = re.sub(r'[-_]?\d{10}', '', n)
    n = re.sub(r'[-_]?\d{8,}', '', n)
    n = re.sub(r'[-_()（）\[\]]+', '', n)
    for pat in [r'眼科学?期末.*', r'眼科学?出题.*', r'眼科学?考题.*',
                r'眼科学?试题.*', r'眼科学?作业.*', r'眼科题目.*',
                r'眼科命题.*', r'眼科出题.*', r'眼科考题.*', r'眼科出題.*',
                r'眼科题.*', r'眼科$', r'眼科见习复习题']:
        n = re.sub(pat, '', n)
    return n.strip() if n.strip() and len(n.strip()) >= 2 else fname[:20]

# Manual author cleanups
AUTHOR_FIXES = {
    '谭饮骏-2110122901': '谭饮骏',
    '韦鹏昊眼科学综合试题': '韦鹏昊',
    '刘霁影45题1': '刘霁影',
}

def parse_questions(text):
    """Universal parser covering every observed format."""
    questions = []
    lines = [l.strip() for l in text.split('\n') if l.strip()]

    # ====== PHASE 0: Find answer key block (韦鹏昊 style) ======
    answer_key = {}
    expl_map = {}

    # Answer speed key: "1.B　　2.C"
    ak = re.search(r'答案速查\s*\n(.+?)(?:\n\s*\n|\n(?:二|三|四)[、\.])', text, re.DOTALL)
    if ak:
        for m in re.finditer(r'(\d+)\s*[\.\、]\s*([A-Ea-e])', ak.group(1)):
            answer_key[int(m.group(1))] = m.group(2).upper()

    # Detailed answer blocks: "N. 【答案】X\n【解析】..."
    for m in re.finditer(
        r'(\d+)\.\s*【答案】\s*([A-Ea-e])\s*\n【解析】\s*(.+?)(?=\n\d+\.\s*【答案】|\n(?:三|四|五)[、\.]|\Z)',
        text, re.DOTALL):
        qnum = int(m.group(1))
        ans = m.group(2).upper()
        expl = m.group(3).strip()
        expl_map[qnum] = (ans, expl)

    # ====== PHASE 1: Parse questions line by line ======
    i = 0
    current_q = None
    current_opts = {}

    while i < len(lines):
        line = lines[i]

        # Skip headers and separators
        if re.match(r'^-{3,}$', line) or re.match(r'^(?:一|二|三|四|五)[、\.]', line):
            i += 1; continue
        if re.match(r'^【第[一二三四五六七八九十]', line):  # 【第一课】
            i += 1; continue

        # ── Question detection ──
        q_match = None
        q_num = None

        # Pattern 1: 【N】 text (周智睿, 战鹏宇)
        m = re.match(r'^【(\d+)】\s*(.+)', line)
        if m: q_match, q_num = m, int(m.group(1))

        # Pattern 2: N. text or N、text or N．text (most common)
        if not q_match:
            m = re.match(r'^(\d+)\s*[\.\、．]\s*(.+)', line)
            if m and len(m.group(2).strip()) >= 3: q_match, q_num = m, int(m.group(1))

        # Pattern 3: 第 N 题 text (周圣淞)
        if not q_match:
            m = re.match(r'^第\s*(\d+)\s*题\s*\n?(.+)', line, re.DOTALL)
            if m: q_match, q_num = m, int(m.group(1))

        if q_match and q_num:
            q_text = q_match.group(2).strip()

            # Skip if looks like option or too short
            if re.match(r'^[A-Ea-e][\.\、\s]', q_text) or len(q_text) < 2:
                i += 1; continue

            # Check for inline answer in question text: "问题？B" or "问题  B"
            # Pattern: question ends with answer letter (with double space or tab)
            inline_ans = re.search(r'[？?]\s*([A-Ea-e])\s*$', q_text)
            if not inline_ans:
                inline_ans = re.search(r'[）\)]\s*([A-Ea-e])\s*$', q_text)
            if not inline_ans:
                inline_ans = re.search(r'\s{2,}([A-Ea-e])\s*$', q_text)
            # Pattern: 问题（）[C] or 问题（）【C】
            if not inline_ans:
                inline_ans = re.search(r'[）\)]\s*[\[【]([A-Ea-e])[\]】]', q_text)

            # Save previous question
            if current_q and current_opts and len(current_opts) >= 2:
                current_q['options'] = current_opts
                questions.append(current_q)

            current_q = {'num': q_num, 'question': re.sub(r'\s*[\[【][A-Ea-e][\]】]?\s*$', '', q_text).strip(),
                         'question_raw': q_text}
            current_opts = {}

            # Record inline answer if found
            if inline_ans:
                current_q['_inline_ans'] = inline_ans.group(1).upper()
                current_q['question'] = re.sub(r'\s{2,}[A-Ea-e]\s*$', '', current_q['question']).strip()
                current_q['question'] = re.sub(r'[\[【][A-Ea-e][\]】]', '', current_q['question']).strip()
                current_q['question'] = re.sub(r'[？?]\s*$', '？', current_q['question']).strip()

            i += 1; continue

        # ── Option detection ──
        # A. text, A、text, A text (with enough content)
        opt_m = re.match(r'^([A-Ea-e])\s*[\.\、．\s]\s*(.+)', line)
        if opt_m and current_q:
            letter = opt_m.group(1).upper()
            txt = opt_m.group(2).strip()
            if len(txt) >= 1 and not re.match(r'^\d', txt[:1]):
                current_opts[letter] = txt
                i += 1; continue

        # ── Answer markers ──
        # "答案：X" "正确答案：X" with optional 解析
        ans_m = re.match(r'^(?:正确)?答案\s*[：:]\s*([A-Ea-e])', line)
        # "【答案】X" (邓敏智, 周圣淞)
        if not ans_m:
            ans_m = re.match(r'^【答案】\s*([A-Ea-e])', line)
        # "【正确答案】 X" (周圣淞)
        if not ans_m:
            ans_m = re.match(r'^【正确答案】\s*([A-Ea-e])', line)

        if ans_m and current_q:
            current_q['answer'] = ans_m.group(1).upper()
            # Extract inline explanation
            rest = re.sub(r'^(?:【正确】)?(?:正确)?答案[：:】]\s*[A-Ea-e]\s*[。，.]?\s*', '', line)
            rest = re.sub(r'^解析[：:]?\s*', '', rest)
            if rest.strip():
                current_q['explanation'] = rest.strip()
            # Collect more explanation lines
            j = i + 1
            extra = []
            while j < len(lines) and j < i + 8:
                nxt = lines[j]
                if re.match(r'^-{3,}$', nxt): j += 1; continue
                if re.match(r'^【(\d+)】', nxt): break
                if re.match(r'^(\d+)\s*[\.\、]', nxt): break
                if re.match(r'^(?:【正确】)?(?:正确)?答案', nxt): break
                if re.match(r'^(?:一|二|三|四|五)[、\.]', nxt): break
                # Check if it's an explanation line
                expl_line = re.sub(r'^【解析】\s*', '', nxt)
                if expl_line:
                    extra.append(expl_line)
                j += 1
            if extra:
                curr = current_q.get('explanation', '')
                current_q['explanation'] = (curr + ' ' + ' '.join(extra)).strip()
            i = j; continue

        # Direct answer on same line: "问题？ 答案：X" (陈宣仁 style)
        inline_direct = re.match(r'^.+\s+答案[：:]\s*([A-Ea-e])\s*$', line)
        if inline_direct and current_q:
            current_q['answer'] = inline_direct.group(1).upper()
            current_q['question'] = re.sub(r'\s+答案[：:]\s*[A-Ea-e]\s*$', '', current_q.get('question', '')).strip()
            i += 1; continue

        # ── Question text continuation ──
        if current_q and not current_opts and len(line) > 3:
            current_q['question'] += ' ' + line
            i += 1; continue

        i += 1

    # Save last question
    if current_q and current_opts and len(current_opts) >= 2:
        current_q['options'] = current_opts
        questions.append(current_q)

    # ====== PHASE 2: Apply answers from various sources ======
    for q in questions:
        qnum = q.get('num', 0)

        # Priority 1: Explicit answer marker found during parsing
        if q.get('answer'):
            pass
        # Priority 2: Inline answer in question text (赵栩禾, 罗智谦, 宋娟范, 张智翔 etc.)
        elif q.get('_inline_ans'):
            q['answer'] = q['_inline_ans']
        # Priority 3: Answer key from end of document
        elif qnum in answer_key:
            q['answer'] = answer_key[qnum]
        # Priority 4: Detailed explanation block
        elif qnum in expl_map:
            q['answer'], q['explanation'] = expl_map[qnum]

        # Clean up
        if '_inline_ans' in q:
            del q['_inline_ans']
        if 'question_raw' in q:
            del q['question_raw']

    return [q for q in questions if q.get('answer') and len(q.get('options', {})) >= 2]


# ====== MAIN ======
all_qs = []

for group in GROUPS:
    gpath = os.path.join(BASE, group)
    if not os.path.exists(gpath): continue
    for fname in sorted(os.listdir(gpath)):
        if fname.startswith('.'): continue
        fpath = os.path.join(gpath, fname)
        try:
            text = extract_docx(fpath) if fname.endswith('.docx') else extract_pdf(fpath)
        except Exception as e:
            print(f"  SKIP {fname}: {e}")
            continue

        author = extract_author(fname)
        author = AUTHOR_FIXES.get(author, author)
        qs = parse_questions(text)
        for q in qs:
            q['author'] = author
            q['group'] = group
            q['filename'] = fname

        print(f"{os.path.basename(fname):<45s} -> {len(qs):3d} Q  [{author}]")
        all_qs.extend(qs)

# Dedup
seen = set()
uniq = []
for q in all_qs:
    key = q['question'][:80].strip()
    if key not in seen:
        seen.add(key)
        uniq.append(q)

print(f"\n{'='*60}")
print(f"Total: {len(all_qs)} raw -> {len(uniq)} unique")

outpath = os.path.join(BASE, 'questions.json')
with open(outpath, 'w', encoding='utf-8') as f:
    json.dump(uniq, f, ensure_ascii=False, indent=2)

from collections import Counter
for a, c in Counter(q['author'] for q in uniq).most_common():
    print(f"  {a}: {c}")
print(f"Saved to {outpath}")

#!/usr/bin/env python3
"""V4 — definitive extractor handling every observed format."""
import docx, pdfplumber, json, re, os

BASE = "/Users/moliex/projects/eyetest"
GROUPS = ["眼科出题第一组", "眼科出题第二组", "眼科出题第三组"]

def extract_docx(fp):
    doc = docx.Document(fp)
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            parts.append(" | ".join(c.text.strip() for c in row.cells))
    return "\n".join(parts)

def extract_pdf(fp):
    return "\n".join([p.extract_text() or '' for p in pdfplumber.open(fp).pages])

def clean_author(fname):
    n = re.sub(r'\.(docx?|pdf)$', '', fname)
    n = re.sub(r'[-_]?\d{10}', '', n)
    n = re.sub(r'[-_]?\d{8,}', '', n)
    n = re.sub(r'[-_()（）\[\]]+', '', n)
    for pat in [r'眼科学?期末.*', r'眼科学?出题.*', r'眼科学?考题.*',
                r'眼科学?试题.*', r'眼科学?作业.*', r'眼科题目.*',
                r'眼科命题.*', r'眼科出题.*', r'眼科考题.*', r'眼科出題.*',
                r'眼科题.*', r'眼科$', r'眼科见习复习题']:
        n = re.sub(pat, '', n)
    return n.strip()[:30] if n.strip() else fname[:20]

FIXES = {'谭饮骏-2110122901':'谭饮骏','韦鹏昊眼科学综合试题':'韦鹏昊','刘霁影45题1':'刘霁影',
         '眼科出题-王雨涵.docx':'王雨涵','眼科考题刘子越.docx':'刘子越'}

def parse(text):
    """Master parser."""
    questions = []
    lines = [l.strip() for l in text.split('\n') if l.strip()]

    # ═══ PHASE 0: Find answer sections ═══

    # 0a. Answer key block (韦鹏昊 style: "1.B  2.C  3.B")
    key_from_block = {}
    ak = re.search(r'答案速查\s*\n(.+?)(?:\n\s*\n|\n(?:二|三|四)[、\.])', text, re.DOTALL)
    if ak:
        for m in re.finditer(r'(\d+)\s*[\.\、]\s*([A-Ea-e])', ak.group(1)):
            key_from_block[int(m.group(1))] = m.group(2).upper()

    # 0b. Detailed answer blocks: "N. 【答案】X\n【解析】..."
    expl_map = {}
    for m in re.finditer(
        r'(\d+)\.\s*【答案】\s*([A-Ea-e])\s*\n【解析】\s*(.+?)(?=\n\d+\.\s*【答案】|\n(?:三|四|五)[、\.]|\Z)',
        text, re.DOTALL):
        expl_map[int(m.group(1))] = (m.group(2).upper(), m.group(3).strip())

    # 0c. 陈栐洁 style: "参考答案\n一、单选题...\nC （explanation）\nB （explanation）..."
    ordered_answers = []
    ref_match = re.search(r'参考答案\s*\n一、单选题[^\n]*\n(.+?)(?:\n\n|\n(?:二|三)[、\.])', text, re.DOTALL)
    if ref_match:
        # Each line: "C （基质层占..." or "C" or "C （some text）"
        for line in ref_match.group(1).split('\n'):
            m = re.match(r'^\s*([A-Ea-e])\s*[（(]', line)
            if not m:
                m = re.match(r'^\s*([A-Ea-e])\s*$', line)
            if m:
                ordered_answers.append(m.group(1).upper())

    # ═══ PHASE 1: Line-by-line parsing ═══
    i = 0
    current_q = None
    current_opts = {}
    seq_num = 0  # For unnumbered questions

    while i < len(lines):
        line = lines[i]

        # Skip headers, separators, topic labels
        if (re.match(r'^-{3,}$', line) or
            re.match(r'^(?:一|二|三|四|五|六|七|八|九|十)[、\.\s]', line) or
            re.match(r'^【第[一二三四五六七八九十]', line) or
            re.match(r'^【[^】\d]+】$', line)):  # 【绪论与解剖生理】
            i += 1; continue

        # ── Question start detection ──
        q_match = None; q_num = None; q_text = None; inline_ans = None

        # Format A: 【N】text (周智睿, 战鹏宇)
        m = re.match(r'^【(\d+)】\s*(.+)', line)
        if m and len(m.group(2)) >= 2:
            q_match, q_num, q_text = m, int(m.group(1)), m.group(2)

        # Format B: N. text or N、text (most common)
        if not q_match:
            m = re.match(r'^(\d+)\s*[\.\、．]\s*(.+)', line)
            if m and len(m.group(2)) >= 2:
                q_text_raw = m.group(2)
                q_match, q_num, q_text = m, int(m.group(1)), q_text_raw
                # Detect inline answer: "1. question X" (single space before letter at end)
                am = re.search(r'\s+([A-Ea-e])\s*$', q_text_raw)
                if am:
                    inline_ans = am.group(1).upper()
                    q_text = re.sub(r'\s+[A-Ea-e]\s*$', '', q_text_raw)

        # Format C: 第 N 题 (周圣淞 — question text on next line)
        if not q_match:
            m = re.match(r'^第\s*(\d+)\s*题\s*$', line)
            if m and i+1 < len(lines):
                q_num = int(m.group(1))
                q_text = lines[i+1]
                q_match = m
                i += 1  # skip the question text line

        # Format D: Unnumbered question — when we see options followed by 【答案】or 答案
        # (邓敏智, 陈宣仁, 梁瑞博, 黄婉雯, 邬清杨, 陈栐洁)
        if not q_match and not current_q:
            # Detect if line looks like a question (not an option, not an answer)
            is_option = bool(re.match(r'^[A-Ea-e]\s*[\.\、\s]', line))
            is_answer = bool(re.match(r'^(?:【?正确】?)?答案', line))
            is_numeric = bool(re.match(r'^\d+[\.\、]', line))
            if not is_option and not is_answer and not is_numeric and len(line) > 8:
                # Might be an unnumbered question
                seq_num += 1
                q_num = seq_num
                q_text = line
                q_match = True
                # Detect trailing answer for 赵栩禾 style: "question text  B"
                am = re.search(r'\s{2,}([A-Ea-e])\s*$', line)
                if am:
                    inline_ans = am.group(1).upper()
                    q_text = re.sub(r'\s{2,}[A-Ea-e]\s*$', '', line)

        if q_match and q_num:
            # Save previous
            if current_q and current_opts and len(current_opts) >= 2:
                current_q['options'] = dict(current_opts)
                questions.append(current_q)

            current_q = {'num': q_num, 'question': q_text.strip()}
            current_opts = {}
            if inline_ans:
                current_q['_ia'] = inline_ans
            i += 1; continue

        # ── Option lines: A. A、A (text) ──
        opt_m = re.match(r'^([A-Ea-e])\s*[\.\、．\s]\s*(.+)', line)
        if opt_m and current_q:
            letter = opt_m.group(1).upper()
            txt = opt_m.group(2).strip()
            if len(txt) >= 1:
                current_opts[letter] = txt
                i += 1; continue

        # ── Answer markers ──
        # "答案：X" / "正确答案：X" / "【答案】X" / "【正确答案】 X"
        ans_m = re.match(r'^(?:【?正确】?)?(?:正确)?答案[：:】]\s*([A-Ea-e])', line)
        if not ans_m:
            ans_m = re.match(r'^【答案】\s*([A-Ea-e])', line)

        if ans_m and current_q:
            current_q['answer'] = ans_m.group(1).upper()
            # Inline explanation
            rest = re.sub(r'^(?:【?正确】?)?(?:正确)?答案[：:】]\s*[A-Ea-e]\s*[。，.]?\s*', '', line)
            rest = re.sub(r'^解析[：:]?\s*', '', rest)
            if rest.strip(): current_q['explanation'] = rest.strip()
            # Collect continuation
            j = i + 1; extra = []
            while j < len(lines) and j < i + 10:
                nxt = lines[j]
                if re.match(r'^-{3,}$', nxt): j += 1; continue
                if re.match(r'^【(\d+)】', nxt): break
                if re.match(r'^(\d+)\s*[\.\、]', nxt): break
                if re.match(r'^(?:【?正确】?)?(?:正确)?答案', nxt): break
                if re.match(r'^(?:一|二|三|四|五)[、\.]', nxt): break
                if re.match(r'^第\s*\d+\s*题', nxt): break
                expl_line = re.sub(r'^【解析】\s*', '', nxt)
                if expl_line: extra.append(expl_line)
                j += 1
            if extra:
                curr = current_q.get('explanation', '')
                current_q['explanation'] = (curr + ' ' + ' '.join(extra)).strip()
            i = j; continue

        # ── 陈宣仁 style: "答案：B" on line AFTER question ──
        ans_m2 = re.match(r'^答案[：:]\s*([A-Ea-e])\s*$', line)
        if ans_m2 and current_q:
            current_q['answer'] = ans_m2.group(1).upper()
            i += 1; continue

        # ── Question text continuation ──
        if current_q and not current_opts and len(line) > 3:
            # Check if this continuation has inline answer
            am = re.search(r'\s+([A-Ea-e])\s*$', line)
            if am and not re.match(r'^[A-Ea-e]', line[:1]):
                current_q['_ia'] = am.group(1).upper()
                line = re.sub(r'\s+[A-Ea-e]\s*$', '', line)
            current_q['question'] += ' ' + line
            i += 1; continue

        i += 1

    # Save last
    if current_q and current_opts and len(current_opts) >= 2:
        current_q['options'] = dict(current_opts)
        questions.append(current_q)

    # ═══ PHASE 2: Apply answers ═══
    # For 陈栐洁 style (ordered answers), assign sequentially
    if ordered_answers:
        # Questions from this file should be in order
        for idx, q in enumerate(questions):
            if not q.get('answer') and idx < len(ordered_answers):
                q['answer'] = ordered_answers[idx]

    for q in questions:
        qnum = q.get('num', 0)
        # Priority: explicit answer > inline answer > answer block > expl map
        if not q.get('answer') and q.get('_ia'):
            q['answer'] = q['_ia']
        if not q.get('answer') and qnum in key_from_block:
            q['answer'] = key_from_block[qnum]
        if not q.get('answer') and qnum in expl_map:
            q['answer'], q['explanation'] = expl_map[qnum]

        # Cleanup
        q.pop('_ia', None)

    return [q for q in questions if q.get('answer') and len(q.get('options', {})) >= 2]


# ═══ MAIN ═══
all_qs = []
for group in GROUPS:
    gpath = os.path.join(BASE, group)
    if not os.path.exists(gpath): continue
    for fname in sorted(os.listdir(gpath)):
        if fname.startswith('.'): continue
        fpath = os.path.join(gpath, fname)
        try:
            text = extract_docx(fpath) if fname.endswith('.docx') else extract_pdf(fpath)
        except Exception as e:
            print(f"  SKIP {fname}: {e}"); continue

        author = FIXES.get(clean_author(fname), clean_author(fname))
        qs = parse(text)
        for q in qs:
            q['author'] = author; q['group'] = group; q['filename'] = fname
        print(f"{fname:<45s} -> {len(qs):3d}  [{author}]")
        all_qs.extend(qs)

# Dedup
seen = set(); uniq = []
for q in all_qs:
    key = q['question'][:80].strip()
    if key not in seen: seen.add(key); uniq.append(q)

print(f"\n{'='*60}")
print(f"Total: {len(all_qs)} raw -> {len(uniq)} unique")

outpath = os.path.join(BASE, 'questions.json')
with open(outpath, 'w', encoding='utf-8') as f:
    json.dump(uniq, f, ensure_ascii=False, indent=2)

from collections import Counter
for a, c in Counter(q['author'] for q in uniq).most_common():
    print(f"  {a}: {c}")
print(f"Saved to {outpath}")
